"""Extract plain text from common outline/document formats."""
from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path


ALLOWED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".text",
    ".docx", ".doc",  # .doc: best-effort
    ".pdf",
    ".rtf",
    ".csv", ".tsv",
    ".json", ".jsonl",
    ".html", ".htm",
    ".xml",
    ".log",
}

ALLOWED_CONTENT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/x-markdown",
    "text/csv",
    "text/tab-separated-values",
    "text/html",
    "text/xml",
    "text/rtf",
    "application/rtf",
    "application/json",
    "application/xml",
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
    "",
    None,
}


def _decode_bytes(content: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "big5", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _strip_html(html: str) -> str:
    text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", html)
    text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise ValueError("服务器未安装 python-docx，无法解析 .docx") from e
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        raise ValueError("服务器未安装 pypdf，无法解析 .pdf") from e
    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        t = t.strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts).strip()


def _extract_rtf(content: bytes) -> str:
    text = _decode_bytes(content)
    # Best-effort RTF strip without extra dependency
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_text(content: bytes, filename: str | None = None, content_type: str | None = None) -> str:
    """Return plain text from uploaded bytes based on extension/mime."""
    name = (filename or "upload.txt").lower()
    ext = Path(name).suffix

    if ext and ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型: {ext}。支持: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if content_type and content_type not in ALLOWED_CONTENT_TYPES:
        # still allow if extension is known
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(f"不支持的 Content-Type: {content_type}")

    if ext == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = _extract_docx(content)
    elif ext == ".pdf" or content_type == "application/pdf":
        text = _extract_pdf(content)
    elif ext == ".rtf" or content_type in {"text/rtf", "application/rtf"}:
        text = _extract_rtf(content)
    elif ext in {".html", ".htm"}:
        text = _strip_html(_decode_bytes(content))
    elif ext == ".doc":
        # Legacy .doc is binary OLE — try decode as last resort, warn in content
        text = _decode_bytes(content)
        if "\x00" in text[:200]:
            raise ValueError("旧版 .doc 二进制格式无法可靠解析，请另存为 .docx 或 .txt")
    elif ext in {".json", ".jsonl"}:
        text = _decode_bytes(content)
    else:
        text = _decode_bytes(content)

    text = text.strip()
    if not text:
        raise ValueError("未能从文件中提取到文本内容")
    return text
