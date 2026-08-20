from pathlib import Path

from docx import Document

from app.engine.document_blocks import extract_file
from app.engine.import_pipeline import _sanitize_chunks
from app.engine.import_chunker import chunk_document_blocks, coverage_report
from app.document_parsers.docx_parser import parse_docx
from app.document_parsers.plain_text import parse_text


def test_sanitizer_chunks_cover_all_review_blocks():
    blocks = [
        {"block_id": f"review-{i}", "type": "paragraph", "text": f"review-marker-{i} " + ("y" * 500), "section_path": []}
        for i in range(120)
    ]
    chunks = _sanitize_chunks(blocks)
    report = coverage_report(blocks, chunks)
    assert report["coverage_pct"] == 100.0
    assert report["unprocessed_blocks"] == 0
    assert len(chunks) > 1


def test_chunker_covers_100k_document_without_truncation():
    blocks = [
        {"block_id": f"b-{i}", "type": "paragraph", "text": f"source-marker-{i} " + ("x" * 120), "section_path": [f"section-{i // 100}"]}
        for i in range(800)
    ]
    chunks = chunk_document_blocks(blocks)
    report = coverage_report(blocks, chunks)
    assert report["coverage_pct"] == 100.0
    assert report["unprocessed_blocks"] == 0
    assert len(chunks) > 10
    assert "source-marker-799" in "\n".join(chunk.text for chunk in chunks)
    assert all(chunk.char_count <= 12_000 for chunk in chunks)


def test_chunker_keeps_long_single_block_and_tail():
    blocks = [
        {"block_id": "long", "type": "paragraph", "text": "A" * 30_000, "section_path": []},
        {"block_id": "tail", "type": "paragraph", "text": "TAIL-MARKER", "section_path": ["tail"]},
    ]
    chunks = chunk_document_blocks(blocks)
    report = coverage_report(blocks, chunks)
    assert report["coverage_pct"] == 100.0
    assert "TAIL-MARKER" in chunks[-1].text
    assert any("long" in chunk.block_ids for chunk in chunks)


def test_docx_parser_preserves_table_between_paragraphs(tmp_path: Path):
    path = tmp_path / "plan.docx"
    doc = Document()
    doc.add_paragraph("Before")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.add_paragraph("After")
    doc.save(str(path))

    blocks = parse_docx(path, "doc")
    texts = [block.text for block in blocks]
    assert texts[0] == "Before"
    assert "A | B" in texts
    assert "C | D" in texts
    assert texts[-1] == "After"
    assert all(block.source_locator for block in blocks)


def test_extract_file_does_not_turn_unsupported_binary_into_text(tmp_path: Path):
    path = tmp_path / "payload.bin"
    path.write_bytes(b"\x00\xffnot-a-document")
    try:
        extract_file(path, path.name, "doc")
    except ValueError as exc:
        assert "unsupported document format" in str(exc)
    else:
        raise AssertionError("unsupported binary must be rejected")


def test_plain_text_keeps_full_content_and_source_refs():
    blocks = parse_text("# H\nfirst\nsecond", "doc", source_name="plan.txt")
    assert [block.text for block in blocks] == ["H", "first", "second"]
    assert blocks[1].section_path == ["H"]
    assert blocks[1].source_locator["line"] == 2
